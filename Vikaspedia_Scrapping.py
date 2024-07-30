# To Run This code remember to install this packages
# pip install requests
# pip install beautifulsoup4
# pip install python-docx


import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# Function to fetch and parse the webpage
def fetch_webpage(url):
    response = requests.get(url)
    response.encoding = 'utf-8'  # Ensure the response encoding is set to UTF-8 for Hindi content
    return BeautifulSoup(response.text, 'html.parser')

# Function to remove footer content
def remove_footer(soup):
    footers = soup.find_all(class_='footer')  # Adjust this to match the actual footer class/ID
    for footer in footers:
        footer.decompose()  # Remove the footer from the soup

# Function to extract content from the webpage
def extract_content(soup):
    content = []
    for tag in soup.find_all(['h3', 'p', 'ol', 'li', 'table']):
        if tag.name == 'h3':
            content.append(('h3', tag.get_text(strip=True)))
        elif tag.name == 'p':
            content.append(('p', tag.get_text(strip=True)))
        elif tag.name == 'li':
            content.append(('li', tag.get_text(strip=True)))
        elif tag.name == 'table':
            rows = []
            for tr in tag.find_all('tr'):
                cells = [td.get_text(strip=True) for td in tr.find_all('td')]
                rows.append(cells)
            content.append(('table', rows))
    return content

# Function to create a DOCX file and add the extracted content
def create_docx(content, output_file):
    doc = Document()
    
    for item in content:
        if item[0] == 'h3':
            p = doc.add_heading(level=3)
            run = p.add_run(item[1])
            run.font.size = Pt(14)
            run.bold = True
        elif item[0] == 'p':
            p = doc.add_paragraph(item[1])
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        elif item[0] == 'li':
            doc.add_paragraph(item[1], style='List Bullet')
        elif item[0] == 'table':
            if not item[1]:  # Skip empty tables
                continue
            max_cols = max(len(row) for row in item[1])
            table = doc.add_table(rows=0, cols=max_cols)
            for row in item[1]:
                tr = table.add_row().cells
                for i, cell in enumerate(row):
                    if i < len(tr):  # Ensure index is within the number of columns
                        tr[i].text = cell
    
    doc.save(output_file)

# Main function to process the URL and create a DOCX
def main(url, output_file):
    soup = fetch_webpage(url)
    remove_footer(soup)  # Remove footer content
    content = extract_content(soup)
    create_docx(content, output_file)
    print(f"Content extracted and saved to {output_file}")

# Replace 'your_webpage_url' with the actual URL and 'output.docx' with the desired output file name
main('Vikaspedia_Website_URL', 'File_Name.docx')
